import ttkbootstrap as tb
from ttkbootstrap.constants import *
from ttkbootstrap.scrolled import ScrolledText
from ttkbootstrap.dialogs import Messagebox
import threading
import logging
from concurrent.futures import ThreadPoolExecutor
import os
import time
from translator import BaiduTranslator, TextFormatter
import tkinter as tk
import queue
import pyautogui
from collections import OrderedDict
import weakref
from functools import lru_cache
class BaseUIComponent:
    def __init__(self, parent, settings_manager):
        self.parent = parent
        self.settings_manager = settings_manager

    def _create_frame(self, text):
        frame = tb.Frame(self.parent)
        self.parent.add(frame, text=text)
        return frame

    def _create_labeled_frame(self, parent, text, padding=10, bootstyle=INFO):
        return tb.LabelFrame(parent, text=text, padding=padding, bootstyle=bootstyle)

    def _create_button(self, parent, text, command, bootstyle=PRIMARY, width=10):
        return tb.Button(parent, text=text, command=command, bootstyle=bootstyle, width=width)
class LanguageMapper:
    LANG_MAP = {
        '自动检测': 'auto',
        '中文': 'zh',
        '英语': 'en',
        '日语': 'ja',
        '韩语': 'ko',
        '法语': 'fr',
        '德语': 'de',
        '俄语': 'ru',
        '西班牙语': 'es'
    }

    @classmethod
    def get_lang_code(cls, lang_name):
        return cls.LANG_MAP.get(lang_name, 'auto')
class ErrorHandler:
    @staticmethod
    def handle_error(error, operation="操作"):
        error_msg = f"{operation}失败: {str(error)}"
        logging.error(error_msg)
        Messagebox.show_error("错误", error_msg)
class ConfigManager:
    def __init__(self, settings_manager):
        self.settings_manager = settings_manager
        
    def save_config(self, appid, appkey, shortcuts, theme, source_lang, target_lang):
        try:
            if self.settings_manager.save_all_config(appid, appkey, shortcuts, theme, source_lang, target_lang):
                return True
            return False
        except Exception as e:
            ErrorHandler.handle_error(e, "保存配置")
            return False
            
    def load_config(self):
        try:
            return self.settings_manager.load_config()
        except Exception as e:
            ErrorHandler.handle_error(e, "加载配置")
            return None, None
class StatsManager:
    def __init__(self, settings_manager):
        self.settings_manager = settings_manager

    def update_stats(self, source_text):
        stats = self.settings_manager.load_translation_stats()
        stats['total_translations'] = int(stats.get('total_translations', 0)) + 1
        stats['total_characters'] = int(stats.get('total_characters', 0)) + len(source_text)
        stats['daily_translations'] = int(stats.get('daily_translations', 0)) + 1
        stats['daily_characters'] = int(stats.get('daily_characters', 0)) + len(source_text)
        self.settings_manager.save_translation_stats(stats)
        return stats
        
    def get_stats(self):
        """获取统计数据"""
        return self.settings_manager.load_translation_stats()
class TitleBarManager:
    def __init__(self, root):
        self.root = root
        self.is_maximized = False
        self.normal_geometry = None
        self.title_bar = None
        self.max_btn = None
        
    def setup(self):
        """设置标题栏"""
        # 创建主容器
        self.main_frame = tb.Frame(self.root, bootstyle=PRIMARY)
        self.main_frame.place(x=0, y=0, relwidth=1, relheight=1)
        
        # 创建标题栏
        self.title_bar = tb.Frame(self.main_frame, bootstyle=PRIMARY)
        self.title_bar.pack(fill="x")
        
        # 创建标题和控制按钮容器
        content_frame = tb.Frame(self.title_bar, bootstyle=PRIMARY)
        content_frame.pack(fill="x", expand=True)
        
        # 创建标题和控制按钮
        title_label = tb.Label(content_frame, text="翻译工具", 
                            font=('微软雅黑', 10, 'bold'),
                            bootstyle="primary-inverse.TLabel")
        title_label.pack(side="left", padx=5)
        
        # 创建按钮容器，确保右对齐
        button_frame = tb.Frame(content_frame, bootstyle=PRIMARY)
        button_frame.pack(side="right", fill="y", padx=0)
        
        # 添加控制按钮，调整样式
        min_btn = tb.Button(button_frame, text="─", width=3,
                        bootstyle="primary.TButton",
                        command=self.root.iconify)
        min_btn.pack(side="left", padx=(0,1))
        
        self.max_btn = tb.Button(button_frame, text="□", width=3,
                            bootstyle="primary.TButton",
                            command=self.toggle_maximize)
        self.max_btn.pack(side="left", padx=1)
        
        close_btn = tb.Button(button_frame, text="✕", width=3,
                            bootstyle="danger.TButton",
                            command=self.root.quit)
        close_btn.pack(side="left", padx=(1,0))

        # 绑定拖动事件到整个标题栏
        self.title_bar.bind('<Button-1>', self.start_move)
        self.title_bar.bind('<B1-Motion>', self.on_move)
        content_frame.bind('<Button-1>', self.start_move)
        content_frame.bind('<B1-Motion>', self.on_move)
        title_label.bind('<Button-1>', self.start_move)
        title_label.bind('<B1-Motion>', self.on_move)


    def update_style(self, theme):
        """更新标题栏样式"""
        try:
            if theme == '黑夜':
                self.title_bar.configure(bootstyle='dark')
                # 更新所有按钮样式
                for child in self.title_bar.winfo_children():
                    if isinstance(child, tb.Frame):
                        for btn in child.winfo_children():
                            if isinstance(btn, tb.Button):
                                if btn.cget('text') in ['─', '□']:  # 最小化和最大化按钮
                                    btn.configure(bootstyle='dark.TButton')
                                elif btn.cget('text') == '✕':  # 关闭按钮
                                    btn.configure(bootstyle='danger.TButton')  # 保持红色
            else:
                self.title_bar.configure(bootstyle='primary')
                # 更新所有按钮样式
                for child in self.title_bar.winfo_children():
                    if isinstance(child, tb.Frame):
                        for btn in child.winfo_children():
                            if isinstance(btn, tb.Button):
                                if btn.cget('text') in ['─', '□']:  # 最小化和最大化按钮
                                    btn.configure(bootstyle='primary.TButton')
                                elif btn.cget('text') == '✕':  # 关闭按钮
                                    btn.configure(bootstyle='danger.TButton')  # 保持红色
        except Exception as e:
            logging.error(f"更新标题栏样式失败: {str(e)}")

    def start_move(self, event):
        """开始移动窗口"""
        self.x = event.x
        self.y = event.y

    def on_move(self, event):
        """移动窗口"""
        deltax = event.x - self.x
        deltay = event.y - self.y
        x = self.root.winfo_x() + deltax
        y = self.root.winfo_y() + deltay
        self.root.geometry(f"+{x}+{y}")

    def toggle_maximize(self):
        """切换窗口最大化状态"""
        if not self.is_maximized:
            self.normal_geometry = self.root.geometry()
            self.root.geometry(f"{self.root.winfo_screenwidth()}x{self.root.winfo_screenheight()}+0+0")
            self.max_btn.config(text="❐")
            self.is_maximized = True
        else:
            self.root.geometry(self.normal_geometry)
            self.max_btn.config(text="□")
            self.is_maximized = False
class EventBindingMixin:
    def bind_shortcuts(self):
        try:
            shortcuts = self.settings_manager.load_shortcuts()
            self.root.unbind_class('Toplevel', '<Control-Return>')
            self.root.unbind_class('Toplevel', '<Control-d>')
            self.root.unbind_class('Toplevel', '<Control-s>')
            
            self.root.bind(shortcuts.get('translate', '<Control-Return>'), lambda e: self.translate())
            self.root.bind(shortcuts.get('clear', '<Control-d>'), lambda e: self.clear_text())
            self.root.bind(shortcuts.get('capture', '<Control-s>'), lambda e: self.capture_translate())
            
        except Exception as e:
            ErrorHandler.handle_error(e, "绑定快捷键")
class TranslateTabManager(BaseUIComponent):
    def __init__(self, notebook, settings_manager):
        super().__init__(notebook, settings_manager)
        self.notebook = notebook
        self._init_variables()
    def _init_variables(self):
        self.source_lang = None
        self.target_lang = None
        self.source_text = None
        self.target_text = None
        self.buttons = {}
    def _create_button(self, parent, text, command, bootstyle=PRIMARY, width=10):
        return super()._create_button(parent, text, command, bootstyle, width)     
    def setup(self):
        """设置翻译标签页"""
        translate_frame = tb.Frame(self.notebook)
        self.notebook.add(translate_frame, text="📸 截图翻译")

        translate_frame.rowconfigure(1, weight=1)
        translate_frame.columnconfigure(0, weight=1)

        self._create_toolbar(translate_frame)
        self._create_text_areas(translate_frame)
        self._create_bottom_toolbar(translate_frame)

    def _create_toolbar(self, parent):
        """创建工具栏"""
        toolbar = tb.Frame(parent, bootstyle=SECONDARY)
        toolbar.grid(row=0, column=0, sticky="ew", padx=10, pady=(10, 5))

        lang_frame = tb.LabelFrame(toolbar, text="语言设置", padding=6, bootstyle=INFO)
        lang_frame.pack(side=TOP, padx=0, pady=0, fill=X)

        lang_grid = tb.Frame(lang_frame)
        lang_grid.pack(fill=X)
        lang_grid.columnconfigure(1, weight=1)
        lang_grid.columnconfigure(3, weight=1)

        tb.Label(lang_grid, text="源语言:", font=('微软雅黑', 9)).grid(row=0, column=0, padx=(5,2), pady=2, sticky=E)
        self.source_lang = tb.Combobox(lang_grid, width=12, state="readonly", bootstyle=INFO)
        self.source_lang['values'] = LanguageMapper.LANG_MAP.keys()
        self.source_lang.set('自动检测')
        self.source_lang.grid(row=0, column=1, padx=(0,10), pady=2, sticky=W+E)

        tb.Label(lang_grid, text="目标语言:", font=('微软雅黑', 9)).grid(row=0, column=2, padx=(5,2), pady=2, sticky=E)
        self.target_lang = tb.Combobox(lang_grid, width=12, state="readonly", bootstyle=INFO)
        self.target_lang['values'] = list(LanguageMapper.LANG_MAP.keys())[1:]
        self.target_lang.set('英语')
        self.target_lang.grid(row=0, column=3, padx=(0,5), pady=2, sticky=W+E)

    def _create_text_areas(self, parent):
        """创建文本区域"""
        text_container = tb.Frame(parent)
        text_container.grid(row=1, column=0, sticky="nsew", padx=10, pady=5)
        text_container.rowconfigure(0, weight=1)
        text_container.columnconfigure(0, weight=1)

        paned_window = tb.PanedWindow(text_container, orient=VERTICAL)
        paned_window.pack(fill=BOTH, expand=True)

        source_frame = tb.LabelFrame(paned_window, text="源文本", padding=6, bootstyle=PRIMARY)
        paned_window.add(source_frame, weight=1)

        self.source_text = ScrolledText(source_frame, wrap="word", height=5, font=('微软雅黑', 10))
        self.source_text.pack(padx=4, pady=4, fill=BOTH, expand=True)

        target_frame = tb.LabelFrame(paned_window, text="翻译结果", padding=6, bootstyle=PRIMARY)
        paned_window.add(target_frame, weight=2)

        self.target_text = ScrolledText(target_frame, wrap="word", height=7, font=('微软雅黑', 10))
        self.target_text.pack(padx=4, pady=4, fill=BOTH, expand=True)

    def _create_bottom_toolbar(self, parent):
        """创建底部工具栏"""
        bottom_toolbar = tb.Frame(parent, bootstyle=SECONDARY)
        bottom_toolbar.grid(row=2, column=0, sticky="ew", padx=10, pady=(5, 10))

        button_frame = tb.LabelFrame(bottom_toolbar, text="操作", padding=6, bootstyle=INFO)
        button_frame.pack(side=TOP, padx=0, pady=0, fill=X)

        button_grid = tb.Frame(button_frame)
        button_grid.pack(fill=X)

        # 创建按钮并添加工具提示
        self.translate_btn = tb.Button(button_grid, text="🔤 翻译", 
                                    bootstyle=PRIMARY, width=10)
        self.translate_btn.grid(row=0, column=0, padx=6, pady=2, sticky="ew")
        Tooltip(self.translate_btn, "翻译选中的文本 (Ctrl+Enter)")

        self.clear_btn = tb.Button(button_grid, text="🗑️ 清空",
                            bootstyle=WARNING, width=10)
        self.clear_btn.grid(row=0, column=1, padx=6, pady=2, sticky="ew")
        Tooltip(self.clear_btn, "清空所有文本 (Ctrl+D)")

        self.capture_btn = tb.Button(button_grid, text="📷 截图翻译",
                            bootstyle=INFO, width=10)
        self.capture_btn.grid(row=0, column=2, padx=6, pady=2, sticky="ew")
        Tooltip(self.capture_btn, "截图并翻译 (Ctrl+S)")

        self.speak_btn = tb.Button(button_grid, text="🔊 朗读",
                            bootstyle=SUCCESS, width=10)
        self.speak_btn.grid(row=0, column=3, padx=6, pady=2, sticky="ew")
        Tooltip(self.speak_btn, "朗读翻译结果")

        self.export_btn = tb.Button(button_grid, text="📄 导出",
                                bootstyle=INFO, width=10)
        self.export_btn.grid(row=0, column=4, padx=6, pady=2, sticky="ew")
        Tooltip(self.export_btn, "导出翻译结果")

        button_grid.columnconfigure(0, weight=1)
        button_grid.columnconfigure(1, weight=1)
        button_grid.columnconfigure(2, weight=1)
        button_grid.columnconfigure(3, weight=1)
        button_grid.columnconfigure(4, weight=1)

class ConfigTabManager(BaseUIComponent):
    def __init__(self, notebook, settings_manager):
        super().__init__(notebook, settings_manager)
        self.notebook = notebook
        self.settings_manager = settings_manager
        self.theme_var = None
        self.appid_entry = None
        self.appkey_entry = None
        self.translate_shortcut = None
        self.clear_shortcut = None
        self.capture_shortcut = None
        
    def setup(self):
        """设置配置标签页"""
        config_frame = tb.Frame(self.notebook)
        self.notebook.add(config_frame, text="⚙️ 配置")

        # 创建主容器，使用pack布局
        main_container = tb.Frame(config_frame)
        main_container.pack(fill=BOTH, expand=True, padx=10, pady=10)

        # 创建左右两列
        left_column = tb.Frame(main_container)
        left_column.pack(side=LEFT, fill=BOTH, expand=True)

        right_column = tb.Frame(main_container)
        right_column.pack(side=RIGHT, fill=BOTH, padx=(10, 0))

        # 在左列中创建设置区域
        self._create_api_settings(left_column)
        self._create_shortcut_settings(left_column)
        self._create_format_settings(left_column)
        self._create_export_settings(left_column)

        # 在右列中创建设置区域
        self._create_theme_settings(right_column)
        self._create_save_button(right_column)

    def _create_theme_settings(self, parent):
        """创建主题设置区域"""
        # 直接添加主题选择，不使用分组框
        theme_container = tb.Frame(parent)
        theme_container.pack(fill=X, pady=(0, 10))

        tb.Label(theme_container, text="界面主题:", 
                font=('微软雅黑', 10, 'bold')).pack(side="left", padx=5)
        
        self.theme_var = tb.StringVar()
        self.theme_combo = tb.Combobox(theme_container, width=10, state="readonly",
                                    textvariable=self.theme_var, bootstyle=PRIMARY)
        self.theme_combo['values'] = ('白天', '黑夜')
        self.theme_combo.set('白天')
        self.theme_combo.pack(side="left", padx=5)

    def _create_api_settings(self, parent):
        """创建API设置区域"""
        api_frame = tb.LabelFrame(parent, text="百度翻译API", padding=8, bootstyle=INFO)
        api_frame.pack(fill=X, pady=5)

        # 添加提示文本
        hint_label = tb.Label(api_frame, text="请输入百度翻译开放平台的API密钥",
                            font=('微软雅黑', 9),
                            bootstyle=INFO)
        hint_label.pack(pady=(0, 5))

        # API ID输入框
        id_frame = tb.Frame(api_frame)
        id_frame.pack(fill=X, pady=(0, 5))
        tb.Label(id_frame, text="APP ID:", width=10).pack(side=LEFT)
        self.appid_entry = tb.Entry(id_frame, bootstyle=PRIMARY, font=('微软雅黑', 10))
        self.appid_entry.pack(side=LEFT, fill=X, expand=True)

        # API Key输入框
        key_frame = tb.Frame(api_frame)
        key_frame.pack(fill=X)
        tb.Label(key_frame, text="APP Key:", width=10).pack(side=LEFT)
        self.appkey_entry = tb.Entry(key_frame, show="*", bootstyle=PRIMARY, font=('微软雅黑', 10))
        self.appkey_entry.pack(side=LEFT, fill=X, expand=True)

    def _create_shortcut_settings(self, parent):
        """创建快捷键设置区域"""
        shortcuts_frame = tb.LabelFrame(parent, text="快捷键设置", padding=8, bootstyle=INFO)
        shortcuts_frame.pack(fill=X, pady=5)

        # 创建两列布局
        shortcuts_grid = tb.Frame(shortcuts_frame)
        shortcuts_grid.pack(fill=X)
        shortcuts_grid.columnconfigure(0, weight=1)
        shortcuts_grid.columnconfigure(1, weight=1)

        # 第一行：翻译和清空
        tb.Label(shortcuts_grid, text="翻译:", font=('微软雅黑', 9)).grid(row=0, column=0, sticky="w", padx=5, pady=2)
        self.translate_shortcut = tb.Entry(shortcuts_grid, bootstyle=PRIMARY, font=('微软雅黑', 9))
        self.translate_shortcut.grid(row=0, column=1, padx=5, pady=2, sticky="ew")

        tb.Label(shortcuts_grid, text="清空:", font=('微软雅黑', 9)).grid(row=0, column=2, sticky="w", padx=5, pady=2)
        self.clear_shortcut = tb.Entry(shortcuts_grid, bootstyle=PRIMARY, font=('微软雅黑', 9))
        self.clear_shortcut.grid(row=0, column=3, padx=5, pady=2, sticky="ew")

        # 第二行：截图翻译
        tb.Label(shortcuts_grid, text="截图翻译:", font=('微软雅黑', 9)).grid(row=1, column=0, sticky="w", padx=5, pady=2)
        self.capture_shortcut = tb.Entry(shortcuts_grid, bootstyle=PRIMARY, font=('微软雅黑', 9))
        self.capture_shortcut.grid(row=1, column=1, columnspan=3, padx=5, pady=2, sticky="ew")

    def _create_format_settings(self, parent):
        """创建格式化设置区域"""
        format_frame = tb.LabelFrame(parent, text="格式化设置", padding=8, bootstyle=INFO)
        format_frame.pack(fill=X, pady=5)

        # 创建三列布局
        format_grid = tb.Frame(format_frame)
        format_grid.pack(fill=X)
        format_grid.columnconfigure(0, weight=1)
        format_grid.columnconfigure(1, weight=1)
        format_grid.columnconfigure(2, weight=1)

        self.format_var = tb.StringVar(value="none")
        formats = [
            ("无格式", "none"),
            ("保留换行", "keep_newline"),
            ("添加标点", "add_punctuation"),
            ("首字母大写", "capitalize"),
            ("每句换行", "sentence_newline")
        ]
        
        # 第一行：三个选项
        for i, (text, value) in enumerate(formats[:3]):
            tb.Radiobutton(format_grid, text=text, variable=self.format_var, 
                        value=value).grid(row=0, column=i, padx=5, pady=2, sticky="w")
        
        # 第二行：剩余选项
        for i, (text, value) in enumerate(formats[3:], start=3):
            tb.Radiobutton(format_grid, text=text, variable=self.format_var, 
                        value=value).grid(row=1, column=i-3, padx=5, pady=2, sticky="w")

    def _create_export_settings(self, parent):
        """创建导出设置区域"""
        export_frame = tb.LabelFrame(parent, text="导出设置", padding=8, bootstyle=INFO)
        export_frame.pack(fill=X, pady=5)

        self.export_format_var = tb.StringVar(value="txt")
        export_formats = tb.Combobox(export_frame, textvariable=self.export_format_var,
                                    width=12, state="readonly", bootstyle=PRIMARY)
        export_formats['values'] = ('txt', 'docx', 'pdf', 'json')
        export_formats.pack(pady=5)

    def _create_save_button(self, parent):
        """创建保存按钮"""
        # 添加分隔线
        separator = tb.Frame(parent, height=1, bootstyle=SECONDARY)
        separator.pack(fill=X, pady=(20, 10))
        
        self.save_btn = tb.Button(parent, text="💾 保存配置", 
                                bootstyle=SUCCESS, width=15)
        self.save_btn.pack(pady=10)

    def load_shortcuts(self):
        """加载快捷键配置"""
        shortcuts = self.settings_manager.load_shortcuts()
        self.translate_shortcut.delete(0, "end")
        self.translate_shortcut.insert(0, shortcuts.get('translate', '<Control-Return>'))
        self.clear_shortcut.delete(0, "end")
        self.clear_shortcut.insert(0, shortcuts.get('clear', '<Control-d>'))
        self.capture_shortcut.delete(0, "end")
        self.capture_shortcut.insert(0, shortcuts.get('capture', '<Control-s>'))

class AboutTabManager(BaseUIComponent):
    def __init__(self, notebook, settings_manager):
        super().__init__(notebook, settings_manager)
        self.notebook = notebook
        self.settings_manager = settings_manager
        
    def setup(self):
        """设置关于标签页"""
        from version_config import VERSION_INFO
        
        about_frame = tb.Frame(self.notebook)
        self.notebook.add(about_frame, text="ℹ️ 关于")

        about_container = tb.Frame(about_frame)
        about_container.pack(padx=20, pady=20, fill=BOTH, expand=True)

        # 创建Logo和标题区域
        self._create_header(about_container, VERSION_INFO)
        # 创建信息卡片
        self._create_info_card(about_container, VERSION_INFO)
        # 创建统计卡片
        self._create_stats_card(about_container)
        # 创建功能说明
        self._create_features(about_container, VERSION_INFO)

    def _create_header(self, parent, version_info):
        """创建标题区域"""
        logo_frame = tb.Frame(parent)
        logo_frame.pack(fill=X, pady=(0, 20))
        
        title_label = tb.Label(logo_frame, text="翻译工具", 
                            font=('微软雅黑', 24, 'bold'))
        title_label.pack()
        
        version_label = tb.Label(logo_frame, text=f"Version {version_info['version']}", 
                            font=('微软雅黑', 12))
        version_label.pack()

    def _create_info_card(self, parent, version_info):
        """创建信息卡片"""
        info_card = tb.LabelFrame(parent, text="软件信息", padding=20, bootstyle=INFO)
        info_card.pack(fill=X, pady=10)

        # 作者信息
        author_frame = tb.Frame(info_card)
        author_frame.pack(fill=X, pady=5)
        
        tb.Label(author_frame, text="作者：", font=('微软雅黑', 10)).pack(side=LEFT)
        tb.Label(author_frame, text=version_info['author'], font=('微软雅黑', 10, 'bold')).pack(side=LEFT)

        # GitHub链接
        github_frame = tb.Frame(info_card)
        github_frame.pack(fill=X, pady=5)
        
        tb.Label(github_frame, text="项目地址：", font=('微软雅黑', 10)).pack(side=LEFT)
        github_link = tb.Label(github_frame, text=version_info['github'], 
                            font=('微软雅黑', 10, 'bold'),
                            foreground='#4A90E2',
                            cursor='hand2')
        github_link.pack(side=LEFT)
        github_link.bind("<Button-1>", lambda e: self._open_link(version_info['github']))

    def _create_features(self, parent, version_info):
        """创建功能说明"""
        feature_frame = tb.LabelFrame(parent, text="主要功能", padding=20, bootstyle=INFO)
        feature_frame.pack(fill=X, pady=5)

        # 创建功能行容器
        feature_row = tb.Frame(feature_frame)
        feature_row.pack()
        
        # 功能列表
        for i, feature in enumerate(version_info['features']):
            tb.Label(feature_row, text=feature, font=('微软雅黑', 9)).pack(side=LEFT, padx=10)

    def _open_link(self, url):
        """打开链接"""
        import webbrowser
        webbrowser.open(url)
    def _create_stats_card(self, parent):
        """创建统计信息卡片"""
        stats = self.settings_manager.load_translation_stats()
        
        stats_frame = tb.LabelFrame(parent, text="翻译统计", padding=20, bootstyle=INFO)
        stats_frame.pack(fill=X, pady=10)
        
        # 创建统计行
        stats_row = tb.Frame(stats_frame)
        stats_row.pack(fill=X, pady=5)
        
        # 今日统计
        daily_frame = tb.Frame(stats_row)
        daily_frame.pack(side=LEFT, padx=(0, 20))
        
        tb.Label(daily_frame, text="今日翻译：", font=('微软雅黑', 10)).pack(side=LEFT)
        tb.Label(daily_frame, 
                text=f"{stats['daily_translations']}次 ({stats['daily_characters']}字)",
                font=('微软雅黑', 10, 'bold')).pack(side=LEFT)
        
        # 总计统计
        total_frame = tb.Frame(stats_row)
        total_frame.pack(side=LEFT)
        
        tb.Label(total_frame, text="总计翻译：", font=('微软雅黑', 10)).pack(side=LEFT)
        tb.Label(total_frame,
                text=f"{stats['total_translations']}次 ({stats['total_characters']}字)",
                font=('微软雅黑', 10, 'bold')).pack(side=LEFT)
class HistoryTabManager(BaseUIComponent):
    def __init__(self, notebook, settings_manager):
        super().__init__(notebook, settings_manager)
        self.notebook = notebook
        self.settings_manager = settings_manager
        self.history_list = None
        self.clear_btn = None
        self.search_var = None
        self.search_entry = None
        
    def setup(self):
        """设置历史记录标签页"""
        history_frame = tb.Frame(self.notebook)
        self.notebook.add(history_frame, text="📜 历史记录")
        history_frame.pack_propagate(False)
        history_frame.rowconfigure(1, weight=1)
        history_frame.columnconfigure(0, weight=1)

        # 创建工具栏
        toolbar = tb.Frame(history_frame)
        toolbar.grid(row=0, column=0, sticky="ew", padx=10, pady=(10, 5))

        # 搜索框区域
        search_container = tb.Frame(toolbar)
        search_container.pack(side="left", fill="x", expand=True, padx=(0, 10))
        
        tb.Label(search_container, text="搜索:", bootstyle=INFO).pack(side="left", padx=(0, 5))
        self.search_var = tb.StringVar()
        self.search_var.trace('w', self._on_search)
        self.search_entry = tb.Entry(search_container, textvariable=self.search_var, bootstyle=PRIMARY)
        self.search_entry.pack(side="left", fill="x", expand=True)

        # 操作按钮容器
        button_container = tb.Frame(toolbar)
        button_container.pack(side="right")
        # 创建水平排列的按钮框架
        button_frame = tb.Frame(button_container)
        button_frame.pack()
        # 创建导出下拉框
        self.export_var = tb.StringVar()
        self.export_combo = tb.Combobox(button_frame, textvariable=self.export_var, 
                                        width=10, state="readonly", bootstyle=INFO)
        self.export_combo['values'] = ('导出TXT', '导出Word', '导出PDF', '导出JSON')
        self.export_combo.set('导出格式')
        self.export_combo.pack(side="left", padx=5, pady=5)
        self.export_combo.bind('<<ComboboxSelected>>', self._on_export_selected)

        # 清空历史按钮
        self.clear_btn = tb.Button(button_frame, text="清空历史", 
                                bootstyle=DANGER,
                                command=self.clear_history)
        self.clear_btn.pack(side="left", padx=5, pady=5)
        # 创建历史记录列表容器
        list_container = tb.Frame(history_frame)
        list_container.grid(row=1, column=0, sticky="nsew", padx=10, pady=(0, 10))
        list_container.rowconfigure(0, weight=1)
        list_container.columnconfigure(0, weight=1)

        # 创建历史记录列表
        self.history_list = tb.Treeview(list_container, 
                                      columns=('time', 'source', 'target', 'from_lang', 'to_lang'),
                                      show='headings')
        
        # 设置列标题和宽度
        self.history_list.heading('time', text='时间')
        self.history_list.heading('source', text='原文')
        self.history_list.heading('target', text='译文')
        self.history_list.heading('from_lang', text='源语言')
        self.history_list.heading('to_lang', text='目标语言')
        
        # 设置列宽，使用相对宽度
        self.history_list.column('time', width=120, minwidth=100)
        self.history_list.column('source', width=200, minwidth=150)
        self.history_list.column('target', width=200, minwidth=150)
        self.history_list.column('from_lang', width=80, minwidth=60)
        self.history_list.column('to_lang', width=80, minwidth=60)
        
        # 添加滚动条
        v_scrollbar = tb.Scrollbar(list_container, orient=VERTICAL, command=self.history_list.yview)
        h_scrollbar = tb.Scrollbar(list_container, orient=HORIZONTAL, command=self.history_list.xview)
        self.history_list.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)
        
        # 布局滚动条和列表
        self.history_list.grid(row=0, column=0, sticky="nsew")
        v_scrollbar.grid(row=0, column=1, sticky="ns")
        h_scrollbar.grid(row=1, column=0, sticky="ew")
        
        # 绑定双击事件
        self.history_list.bind('<Double-Button-1>', self._show_details)
        
        self.load_history()

    def _show_details(self, event):
        """显示详细信息"""
        selection = self.history_list.selection()
        if not selection:
            return
                
        item = self.history_list.item(selection[0])
        values = item['values']
        
        # 创建详情窗口
        detail_window = tb.Toplevel(self.notebook)
        detail_window.title("翻译详情")
        detail_window.geometry("800x600")
        detail_window.transient(self.notebook)
        detail_window.grab_set()
        
        # 创建主容器
        main_container = tb.Frame(detail_window)
        main_container.pack(fill=BOTH, expand=True, padx=20, pady=20)
        main_container.rowconfigure(1, weight=1)
        main_container.columnconfigure(0, weight=1)

        # 信息栏
        info_frame = tb.Frame(main_container)
        info_frame.grid(row=0, column=0, sticky="ew", pady=(0, 10))
        
        # 时间标签
        tb.Label(info_frame, text=f"翻译时间: {values[0]}", 
                font=('微软雅黑', 10, 'bold')).pack(side="left", padx=(0, 20))
        
        # 语言标签
        tb.Label(info_frame, text=f"翻译: {values[3]} → {values[4]}", 
                font=('微软雅黑', 10)).pack(side="left")

        # 文本区域容器
        text_container = tb.Frame(main_container)
        text_container.grid(row=1, column=0, sticky="nsew")
        text_container.rowconfigure(0, weight=1)
        text_container.rowconfigure(1, weight=1)
        text_container.columnconfigure(0, weight=1)
        text_container.columnconfigure(1, weight=1)

        # 原文区域
        source_frame = tb.LabelFrame(text_container, text="原文", padding=10)
        source_frame.grid(row=0, column=0, sticky="nsew", padx=(0, 5))
        
        source_text = ScrolledText(source_frame, wrap="word", font=('微软雅黑', 10))
        source_text.pack(fill=BOTH, expand=True)
        source_text.text.insert('1.0', values[1])
        source_text.text.configure(state='disabled')

        # 译文区域
        target_frame = tb.LabelFrame(text_container, text="译文", padding=10)
        target_frame.grid(row=0, column=1, sticky="nsew", padx=(5, 0))
        
        target_text = ScrolledText(target_frame, wrap="word", font=('微软雅黑', 10))
        target_text.pack(fill=BOTH, expand=True)
        target_text.text.insert('1.0', values[2])
        target_text.text.configure(state='disabled')

        # 按钮容器
        button_container = tb.Frame(main_container)
        button_container.grid(row=2, column=0, pady=(10, 0))

        # 复制按钮
        copy_source_btn = tb.Button(button_container, text="复制原文", 
                                bootstyle=INFO,
                                command=lambda: self._copy_to_clipboard(values[1]))
        copy_source_btn.pack(side="left", padx=(0, 5))

        copy_target_btn = tb.Button(button_container, text="复制译文", 
                                bootstyle=INFO,
                                command=lambda: self._copy_to_clipboard(values[2]))
        copy_target_btn.pack(side="left", padx=5)

        # 关闭按钮
        close_btn = tb.Button(button_container, text="关闭", 
                            bootstyle=PRIMARY,
                            command=detail_window.destroy)
        close_btn.pack(side="right", padx=(5, 0))

    def _copy_to_clipboard(self, text):
        """复制文本到剪贴板"""
        self.notebook.clipboard_clear()
        self.notebook.clipboard_append(text)

    def _on_search(self, *args):
        """处理搜索事件"""
        search_text = self.search_var.get().lower()
        self.history_list.delete(*self.history_list.get_children())
        
        history = self.settings_manager.load_translation_history()
        for item in history.values():
            if (search_text in item['source_text'].lower() or 
                search_text in item['target_text'].lower()):
                self.history_list.insert('', 'end', values=(
                    item['time'],
                    item['source_text'],
                    item['target_text'],
                    item['from_lang'],
                    item['to_lang']
                ))

    def load_history(self):
        """加载历史记录"""
        self.history_list.delete(*self.history_list.get_children())
        history = self.settings_manager.load_translation_history()
        for item in history.values():
            self.history_list.insert('', 'end', values=(
                item['time'],
                item['source_text'],
                item['target_text'],
                item['from_lang'],
                item['to_lang']
            ))

    def clear_history(self):
        """清空历史记录"""
        if Messagebox.yesno("确认", "确定要清空所有历史记录吗？"):
            self.settings_manager.save_translation_history({})
            self.load_history()
    def _on_export_selected(self, event):
        """处理导出格式选择"""
        try:
            format = self.export_var.get()
            if format == '导出格式':
                return
                
            history = self.settings_manager.load_translation_history()
            if not history:
                Messagebox.show_warning("警告", "没有可导出的历史记录")
                self.export_var.set('导出格式')
                return
                
            if format == '导出TXT':
                self._export_history_txt(history)
            elif format == '导出Word':
                self._export_history_docx(history)
            elif format == '导出PDF':
                self._export_history_pdf(history)
            elif format == '导出JSON':
                self._export_history_json(history)
                
            self.export_var.set('导出格式')
        except Exception as e:
            logging.error(f"导出历史记录失败: {str(e)}")
            Messagebox.show_error("错误", f"导出历史记录失败: {str(e)}")
            self.export_var.set('导出格式')
    def _export_history_txt(self, history):
        """导出历史记录为纯文本"""
        try:
            from tkinter import filedialog
            file_path = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[("Text files", "*.txt"), ("All files", "*.*")]
            )
            if file_path:
                with open(file_path, 'w', encoding='utf-8') as f:
                    for item in history.values():
                        f.write(f"时间: {item['time']}\n")
                        f.write(f"语言: {item['from_lang']} -> {item['to_lang']}\n")
                        f.write(f"原文:\n{item['source_text']}\n")
                        f.write(f"译文:\n{item['target_text']}\n")
                        f.write("-" * 50 + "\n")
                Messagebox.show_info("成功", "导出成功")
        except Exception as e:
            logging.error(f"导出文本失败: {str(e)}")
            Messagebox.show_error("错误", f"导出文本失败: {str(e)}")

    def _export_history_docx(self, history):
        """导出历史记录为Word文档"""
        try:
            from docx import Document
            from tkinter import filedialog
            
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
            )
            if file_path:
                doc = Document()
                doc.add_heading('翻译历史记录', 0)
                
                for item in history.values():
                    doc.add_heading(f"时间: {item['time']}", level=1)
                    doc.add_paragraph(f"语言: {item['from_lang']} -> {item['to_lang']}")
                    doc.add_heading('原文', level=2)
                    doc.add_paragraph(item['source_text'])
                    doc.add_heading('译文', level=2)
                    doc.add_paragraph(item['target_text'])
                    doc.add_paragraph("-" * 50)
                
                doc.save(file_path)
                Messagebox.show_info("成功", "导出成功")
        except ImportError:
            Messagebox.show_error("错误", "导出Word文档需要安装python-docx库")
        except Exception as e:
            logging.error(f"导出Word文档失败: {str(e)}")
            Messagebox.show_error("错误", f"导出Word文档失败: {str(e)}")

    def _export_history_pdf(self, history):
        """导出历史记录为PDF文档"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from tkinter import filedialog
            
            file_path = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
            )
            if file_path:
                c = canvas.Canvas(file_path, pagesize=letter)
                c.setFont("Helvetica", 12)
                
                y = 750
                for item in history.values():
                    if y < 100:  # 换页
                        c.showPage()
                        y = 750
                    
                    c.drawString(100, y, f"时间: {item['time']}")
                    y -= 20
                    c.drawString(100, y, f"语言: {item['from_lang']} -> {item['to_lang']}")
                    y -= 20
                    
                    lines = item['source_text'].split('\n')
                    c.drawString(100, y, "原文:")
                    y -= 20
                    for line in lines:
                        if y < 100:
                            c.showPage()
                            y = 750
                        c.drawString(100, y, line)
                        y -= 20
                    
                    lines = item['target_text'].split('\n')
                    c.drawString(100, y, "译文:")
                    y -= 20
                    for line in lines:
                        if y < 100:
                            c.showPage()
                            y = 750
                        c.drawString(100, y, line)
                        y -= 20
                    
                    y -= 20
                    c.drawString(100, y, "-" * 50)
                    y -= 30
                
                c.save()
                Messagebox.show_info("成功", "导出成功")
        except Exception as e:
            logging.error(f"导出PDF失败: {str(e)}")
            Messagebox.show_error("错误", f"导出PDF失败: {str(e)}")

    def _export_history_json(self, history):
        """导出历史记录为JSON"""
        try:
            import json
            from tkinter import filedialog
            
            file_path = filedialog.asksaveasfilename(
                defaultextension=".json",
                filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
            )
            if file_path:
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(history, f, ensure_ascii=False, indent=2)
                Messagebox.show_info("成功", "导出成功")
        except Exception as e:
            logging.error(f"导出JSON失败: {str(e)}")
            Messagebox.show_error("错误", f"导出JSON失败: {str(e)}")
class UIManager:
    def __init__(self, root, settings_manager):
        self.root = root
        self.settings_manager = settings_manager
        self._translate_lock = threading.Lock()
        self.thread_pool = ThreadPoolExecutor(
            max_workers=5,                                 
            thread_name_prefix="translate_worker"
        )
        self._task_queue = queue.Queue()
        # 初始化管理器
        self.title_bar_manager = TitleBarManager(root)
        self.translate_tab_manager = None
        self.config_tab_manager = None
        self.about_tab_manager = None
        self.translator = None
        self.history_tab_manager = None
        self.stats_manager = StatsManager(settings_manager)
        self.setup_ui()
        
    def setup_ui(self):
        """设置用户界面"""
        try:
            # 创建加载指示器
            self.loading_label = tb.Label(self.root, text="正在加载...", 
                                        font=('微软雅黑', 12))
            self.loading_label.place(relx=0.5, rely=0.5, anchor='center')
            
            # 基本窗口设置
            self.root.overrideredirect(True)
            self.root.style.configure('TNotebook', tabposition='nw')
            self.root.style.configure('TNotebook.Tab', padding=[20, 10])
            
            # 创建主要UI组件
            self.title_bar_manager.setup()
            
            # 修改主容器，去除内边距
            main_container = tb.Frame(self.title_bar_manager.main_frame)
            main_container.pack(fill=BOTH, expand=True)

            # 添加内边距到notebook
            self.notebook = tb.Notebook(main_container, bootstyle=INFO)
            self.notebook.pack(padx=20, pady=20, fill=BOTH, expand=True)
            
            
            # 先创建翻译标签页
            self.translate_tab_manager = TranslateTabManager(self.notebook, self.settings_manager)
            self.translate_tab_manager.setup()
            
            # 延迟加载其他标签页
            self.root.after(100, self._load_other_tabs)
            
            logging.info("界面初始化完成")
        except Exception as e:
            logging.error(f"界面初始化失败: {str(e)}")
            Messagebox.show_error("错误", f"界面初始化失败: {str(e)}")
    def _load_other_tabs(self):
        """延迟加载其他标签页"""
        try:
            # 先创建配置标签页
            self.config_tab_manager = ConfigTabManager(self.notebook, self.settings_manager)
            self.config_tab_manager.setup()
            
            # 等待配置标签页初始化完成后再加载其他标签页
            self.root.after(50, self._load_remaining_tabs)
        except Exception as e:
            logging.error(f"加载配置标签页失败: {str(e)}")
    def _load_remaining_tabs(self):
        """加载剩余标签页"""
        try:
            self.history_tab_manager = HistoryTabManager(self.notebook, self.settings_manager)
            self.about_tab_manager = AboutTabManager(self.notebook, self.settings_manager)
            
            self.history_tab_manager.setup()
            self.about_tab_manager.setup()
            
            # 绑定事件
            self._bind_events()
            
            # 隐藏加载指示器
            if hasattr(self, 'loading_label'):
                self.loading_label.place_forget()
                
            # 通知主应用所有标签页已准备就绪
            self.root.event_generate('<<AllTabsReady>>')
        except Exception as e:
            logging.error(f"加载剩余标签页失败: {str(e)}")
    def _bind_events(self):
        """绑定事件"""
        # 绑定主题切换事件
        self.config_tab_manager.theme_combo.bind("<<ComboboxSelected>>", self.on_theme_change)
        
        # 绑定按钮事件
        self.translate_tab_manager.translate_btn.configure(command=self.translate)
        self.translate_tab_manager.clear_btn.configure(command=self.clear_text)
        self.translate_tab_manager.capture_btn.configure(command=self.capture_translate)
        self.translate_tab_manager.speak_btn.configure(command=self.speak_text)
        self.translate_tab_manager.export_btn.configure(command=self.export_translation)  # 添加导出按钮事件
        self.config_tab_manager.save_btn.configure(command=self.save_config)

        # 绑定快捷键
        self.bind_shortcuts()

    def on_theme_change(self, event=None):
        """处理主题切换"""
        theme = self.config_tab_manager.theme_var.get()
        self.settings_manager.set_theme(theme)
        self.title_bar_manager.update_style(theme)

    def bind_shortcuts(self):
        """绑定快捷键"""
        try:
            # 先解绑所有已存在的快捷键
            self.root.unbind_class('Toplevel', '<Control-Return>')
            self.root.unbind_class('Toplevel', '<Control-d>')
            self.root.unbind_class('Toplevel', '<Control-s>')
            
            # 加载快捷键配置
            shortcuts = self.settings_manager.load_shortcuts()
            
            # 绑定新的快捷键
            self.root.bind(shortcuts.get('translate', '<Control-Return>'), lambda e: self.translate())
            self.root.bind(shortcuts.get('clear', '<Control-d>'), lambda e: self.clear_text())
            self.root.bind(shortcuts.get('capture', '<Control-s>'), lambda e: self.capture_translate())
            
            logging.info("快捷键绑定完成")
        except Exception as e:
            logging.error(f"绑定快捷键失败: {str(e)}")
            Messagebox.show_error("错误", f"绑定快捷键失败: {str(e)}")

    def save_config(self):
        """保存配置"""
        try:
            # 保存API配置
            appid = self.config_tab_manager.appid_entry.get().strip()
            appkey = self.config_tab_manager.appkey_entry.get().strip()
            
            # 保存快捷键配置
            shortcuts = {
                'translate': self.config_tab_manager.translate_shortcut.get(),
                'clear': self.config_tab_manager.clear_shortcut.get(),
                'capture': self.config_tab_manager.capture_shortcut.get()
            }
            
            # 保存主题配置
            theme = self.config_tab_manager.theme_var.get()
            
            # 保存格式化配置
            format_type = self.config_tab_manager.format_var.get()
            
            # 保存导出格式配置
            export_format = self.config_tab_manager.export_format_var.get()
            
            # 保存语言配置
            source_lang = self.translate_tab_manager.source_lang.get()
            target_lang = self.translate_tab_manager.target_lang.get()

            # 保存所有配置
            if self.settings_manager.save_all_config(appid, appkey, shortcuts, theme, source_lang, target_lang, format_type, export_format):
                self.translator = BaiduTranslator(appid, appkey)
                # 设置保存回调
                self.translator.cache.save_callback = self.settings_manager.save_translation_history
                Messagebox.show_info("成功", "配置已保存")
        except Exception as e:
            logging.error(f"保存配置失败: {str(e)}")
            Messagebox.show_error("错误", f"保存配置失败: {str(e)}")

    def load_config(self):
        """加载配置"""
        try:
            # 确保配置标签页已初始化
            if not self.config_tab_manager:
                logging.warning("配置标签页尚未初始化，等待初始化完成...")
                self.root.after(100, self.load_config)  # 延迟重试
                return
                
            # 加载API配置
            appid, appkey = self.settings_manager.load_config()
            if appid and appkey:
                self.config_tab_manager.appid_entry.delete(0, "end")
                self.config_tab_manager.appid_entry.insert(0, appid)
                self.config_tab_manager.appkey_entry.delete(0, "end")
                self.config_tab_manager.appkey_entry.insert(0, appkey)
                self.translator = BaiduTranslator(appid, appkey)
                # 设置保存回调
                self.translator.cache.save_callback = self.settings_manager.save_translation_history
            
            # 加载其他配置
            self._load_remaining_configs()
        except Exception as e:
            logging.error(f"加载配置失败: {str(e)}")
            Messagebox.show_error("错误", f"加载配置失败: {str(e)}")
    def _load_remaining_configs(self):
        """加载剩余配置"""
        try:
            # 确保所有控件都已创建
            if not all([hasattr(self.config_tab_manager, attr) for attr in [
                'translate_shortcut', 'clear_shortcut', 'capture_shortcut',
                'theme_var', 'format_var', 'export_format_var'
            ]]):
                self.root.after(100, self._load_remaining_configs)  # 延迟重试
                return

            # 加载快捷键配置
            self.config_tab_manager.load_shortcuts()
            
            # 加载主题配置
            theme = self.settings_manager.load_theme()
            self.config_tab_manager.theme_var.set(theme)
            self.settings_manager.set_theme(theme)
            
            # 加载格式化配置
            format_type = self.settings_manager.load_format_type()
            self.config_tab_manager.format_var.set(format_type)
            
            # 加载导出格式配置
            export_format = self.settings_manager.load_export_format()
            self.config_tab_manager.export_format_var.set(export_format)
            
            # 加载语言配置
            source_lang, target_lang = self.settings_manager.load_languages()
            if self.translate_tab_manager:
                self.translate_tab_manager.source_lang.set(source_lang)
                self.translate_tab_manager.target_lang.set(target_lang)

            # 加载历史记录到缓存
            if self.translator:
                history = self.settings_manager.load_translation_history()
                self.translator.cache._history = OrderedDict(history)

            logging.info("配置加载完成")
        except Exception as e:
            logging.error(f"加载剩余配置失败: {str(e)}")
            Messagebox.show_error("错误", f"加载配置失败: {str(e)}")
    def translate(self):
        """执行翻译操作"""
        try:
            if not hasattr(self, 'translator') or not self.translator:
                Messagebox.show_error("错误", "请先保存配置")
                return
                
            source_text = self.translate_tab_manager.source_text.get("1.0", "end").strip()
            if not source_text:
                Messagebox.show_warning("警告", "请输入要翻译的文本")
                return
            
            if self._translate_lock.locked():
                Messagebox.show_warning("提示", "正在翻译中，请稍候...")
                return
            
            self._set_controls_state('disabled')
            self.thread_pool.submit(self._translate_thread, source_text)
            
        except Exception as e:
            logging.error(f"翻译操作失败: {str(e)}")
            Messagebox.show_error("错误", f"翻译失败: {str(e)}")
            self._set_controls_state('normal')

    def _translate_thread(self, source_text):
        """翻译线程"""
        try:
            lang_map = {
                '自动检测': 'auto',
                '中文': 'zh',
                '英语': 'en',
                '日语': 'ja',
                '韩语': 'ko',
                '法语': 'fr',
                '德语': 'de',
                '俄语': 'ru',
                '西班牙语': 'es'
            }
            
            from_lang = LanguageMapper.get_lang_code(self.translate_tab_manager.source_lang.get())
            to_lang = LanguageMapper.get_lang_code(self.translate_tab_manager.target_lang.get())
            
            result = self.translator.translate(source_text, from_lang=from_lang, to_lang=to_lang)
            
            self.root.after(0, self._update_result, result)
        except Exception as e:
            ErrorHandler.handle_error(e, "翻译")
        finally:
            self.root.after(0, self._set_controls_state, 'normal')

    def _update_result(self, result):
        """更新翻译结果"""
        try:
            # 获取格式化选项
            format_type = self.translate_tab_manager.format_var.get()
            # 应用格式化
            formatted_result = TextFormatter.format_text(result, format_type)
            
            self.translate_tab_manager.target_text.text.configure(state='normal')
            self.translate_tab_manager.target_text.text.delete("1.0", "end")
            self.translate_tab_manager.target_text.text.insert("1.0", formatted_result)
            self.translate_tab_manager.target_text.text.configure(state='disabled')

            # 更新统计
            source_text = self.translate_tab_manager.source_text.get("1.0", "end").strip()
            self.update_translation_stats(source_text)

            # 添加到历史记录
            from_lang = self.translate_tab_manager.source_lang.get()
            to_lang = self.translate_tab_manager.target_lang.get()
            self.translator.cache.add_to_history(source_text, formatted_result, from_lang, to_lang)
            
            # 更新历史记录显示
            if self.history_tab_manager:
                self.history_tab_manager.load_history()

            logging.info("翻译操作完成")
        except Exception as e:
            logging.error(f"更新翻译结果失败: {str(e)}")
        finally:
            self._set_controls_state('normal')

    def _show_error(self, error_msg):
        """显示错误信息"""
        try:
            error_window = tb.Toplevel(self.root)
            error_window.title("错误提示")
            error_window.geometry("400x200")
            error_window.transient(self.root)
            error_window.grab_set()

            main_frame = tb.Frame(error_window)
            main_frame.pack(fill=BOTH, expand=True, padx=20, pady=20)

            error_icon = tb.Label(main_frame, text="⚠️", font=('微软雅黑', 24))
            error_icon.pack(pady=(0, 10))

            error_label = tb.Label(main_frame, text=error_msg, 
                                font=('微软雅黑', 10))
            error_label.pack(pady=10)

            button_frame = tb.Frame(main_frame)
            button_frame.pack(pady=(10, 0))

            ok_btn = tb.Button(button_frame, text="确定", 
                            bootstyle=PRIMARY,
                            command=error_window.destroy)
            ok_btn.pack()
        except Exception as e:
            logging.error(f"显示错误信息失败: {str(e)}")
        finally:
            self._set_controls_state('normal')

    def _set_controls_state(self, state):
        """设置控件状态"""
        try:
            if state == 'normal':
                self.translate_tab_manager.translate_btn.configure(state='normal')
                self.translate_tab_manager.source_lang.configure(state='readonly')
                self.translate_tab_manager.target_lang.configure(state='readonly')
                self.translate_tab_manager.source_text.text.configure(state='normal')
                self.root.config(cursor="arrow")
            else:
                self.translate_tab_manager.translate_btn.configure(state='disabled')
                self.translate_tab_manager.source_lang.configure(state='disabled')
                self.translate_tab_manager.target_lang.configure(state='disabled')
                self.translate_tab_manager.source_text.text.configure(state='disabled')
                self.root.config(cursor="watch")
        except Exception as e:
            logging.error(f"设置控件状态失败: {str(e)}")

    def clear_text(self):
        """清空文本框"""
        try:
            self.translate_tab_manager.source_text.text.configure(state='normal')
            self.translate_tab_manager.target_text.text.configure(state='normal')
            self.translate_tab_manager.source_text.text.delete("1.0", "end")
            self.translate_tab_manager.target_text.text.delete("1.0", "end")
            self.translate_tab_manager.target_text.text.configure(state='disabled')
            logging.info("清空文本框")
        except Exception as e:
            logging.error(f"清空文本框失败: {str(e)}")
            Messagebox.show_error("错误", f"清空文本框失败: {str(e)}")

    def capture_translate(self):
        """截图翻译"""
        try:
            if not hasattr(self, 'translator') or not self.translator:
                Messagebox.show_error("错误", "请先保存配置")
                return
            self.root.withdraw()
            time.sleep(0.5)
            screenshot = pyautogui.screenshot()
            self._create_selection_window(screenshot)
        except Exception as e:
            logging.error(f"截图翻译失败: {str(e)}")
            Messagebox.show_error("错误", f"截图翻译失败: {str(e)}")
            self.root.deiconify()

    def _minimize_and_capture(self):
        """最小化窗口并执行截图"""
        import pyautogui
        self.root.iconify()
        time.sleep(0.5)
        screenshot = pyautogui.screenshot()
        self._create_selection_window(screenshot)

    def _create_selection_window(self, screenshot):
        """创建选择窗口"""
        import tkinter as tk
        from PIL import Image, ImageTk
        
        selector = tk.Toplevel()
        selector.attributes('-fullscreen', True)
        selector.attributes('-alpha', 0.3)
        selector.configure(background='black')
        
        # 设置窗口始终在最前
        selector.attributes('-topmost', True)
        
        canvas = tk.Canvas(selector, highlightthickness=0)
        canvas.pack(fill=tk.BOTH, expand=True)
        
        # 调整截图大小以适应屏幕
        screen_width = selector.winfo_screenwidth()
        screen_height = selector.winfo_screenheight()
        screenshot = screenshot.resize((screen_width, screen_height))
        
        photo = ImageTk.PhotoImage(screenshot)
        canvas.create_image(0, 0, anchor=tk.NW, image=photo)
        canvas.image = photo
        
        self._setup_selection_events(canvas, selector, screenshot)
        self._add_selection_hints(selector)
        selector.wait_window()

    def _setup_selection_events(self, canvas, selector, screenshot):
        """设置选择事件"""
        start_x = start_y = None
        
        def on_mouse_down(event):
            nonlocal start_x, start_y
            start_x = event.x
            start_y = event.y
            self._create_selection_rect(canvas, start_x, start_y, start_x, start_y)
            
        def on_mouse_drag(event):
            if start_x is not None and start_y is not None:
                self._update_selection_rect(canvas, start_x, start_y, event.x, event.y)
                
        def on_mouse_up(event):
            if start_x is not None and start_y is not None:
                self._process_selection(canvas, selector, screenshot, start_x, start_y, event.x, event.y)
        
        canvas.bind('<Button-1>', on_mouse_down)
        canvas.bind('<B1-Motion>', on_mouse_drag)
        canvas.bind('<ButtonRelease-1>', on_mouse_up)
        selector.bind('<Escape>', lambda e: self._cancel_selection(selector))

    def _create_selection_rect(self, canvas, x1, y1, x2, y2):
        """创建选择框"""
        # 创建外边框
        canvas.create_rectangle(x1-1, y1-1, x2+1, y2+1, outline='white', width=1, tags='selection')
        # 创建主边框
        canvas.create_rectangle(x1, y1, x2, y2, outline='red', width=2, tags='selection')
        # 创建半透明填充
        canvas.create_rectangle(x1, y1, x2, y2, fill='white', stipple='gray25', tags='selection')

    def _update_selection_rect(self, canvas, x1, y1, x2, y2):
        """更新选择框"""
        canvas.delete('selection')
        self._create_selection_rect(canvas, x1, y1, x2, y2)

    def _process_selection(self, canvas, selector, screenshot, x1, y1, x2, y2):
        """处理选择区域"""
        x = min(x1, x2)
        y = min(y1, y2)
        width = abs(x2 - x1)
        height = abs(y2 - y1)
        
        if width < 10 or height < 10:  # 增加最小选择区域
            selector.destroy()
            self.root.deiconify()
            return
        
        # 添加选择确认动画
        canvas.create_rectangle(x, y, x+width, y+height, outline='green', width=3, tags='confirm')
        selector.update()
        time.sleep(0.2)  # 短暂延迟以显示确认效果
        
        selected_area = screenshot.crop((x, y, x + width, y + height))
        
        # 确保data目录存在
        base_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        data_dir = os.path.join(base_path, 'data')
        if not os.path.exists(data_dir):
            os.makedirs(data_dir)
        
        temp_image = os.path.join(data_dir, 'temp_screenshot.png')
        selected_area.save(temp_image)
        
        selector.destroy()
        self._process_ocr_result(temp_image)

    def _process_ocr_result(self, temp_image):
        """处理OCR识别结果"""
        import pytesseract
        from PIL import Image
        
        try:
            text = pytesseract.image_to_string(Image.open(temp_image), lang='chi_sim+eng')
            
            if text.strip():
                self.translate_tab_manager.source_text.text.configure(state='normal')
                self.translate_tab_manager.source_text.text.delete("1.0", "end")
                self.translate_tab_manager.source_text.text.insert("1.0", text.strip())
                self.translate_tab_manager.source_text.text.configure(state='disabled')
                self.translate()
            else:
                Messagebox.show_warning("提示", "未能识别到文本")
        except Exception as e:
            logging.error(f"OCR处理失败: {str(e)}")
            Messagebox.show_error("错误", f"OCR处理失败: {str(e)}")
        finally:
            try:
                os.remove(temp_image)
            except:
                pass
            # 确保窗口恢复显示
            self.root.deiconify()

    def _add_selection_hints(self, selector):
        """添加选择提示"""
        import tkinter as tk
        hint_frame = tk.Frame(selector)
        hint_frame.place(relx=0.5, rely=0.1, anchor='center')
        
        tk.Label(
            hint_frame,
            text="按住鼠标左键并拖动来选择要翻译的区域",
            bg='black',
            fg='white',
            font=('微软雅黑', 12)
        ).pack()

        tk.Label(
            hint_frame,
            text="按 ESC 键取消",
            bg='black',
            fg='gray',
            font=('微软雅黑', 10)
        ).pack()

    def _cancel_selection(self, selector):
        """取消选择"""
        selector.destroy()
        self.root.deiconify()
    def speak_text(self):
        """朗读翻译结果"""
        try:
            if not hasattr(self, 'translator') or not self.translator:
                Messagebox.show_error("错误", "请先保存配置")
                return
                
            text = self.translate_tab_manager.target_text.get("1.0", "end").strip()
            if not text:
                Messagebox.show_warning("警告", "没有可朗读的文本")
                return
                
            # 获取目标语言
            lang_map = {
                '中文': 'zh',
                '英语': 'en',
                '日语': 'ja',
                '韩语': 'ko',
                '法语': 'fr',
                '德语': 'de',
                '俄语': 'ru',
                '西班牙语': 'es'
            }
            target_lang = LanguageMapper.get_lang_code(self.translate_tab_manager.target_lang.get())
            
            self.translator.speak(text, target_lang)
        except Exception as e:
            logging.error(f"朗读失败: {str(e)}")
            Messagebox.show_error("错误", f"朗读失败: {str(e)}")
    def update_translation_stats(self, source_text):
        try:
            self.stats_manager.update_stats(source_text)
            self._update_stats_display(self.stats_manager.get_stats())
        except Exception as e:
            ErrorHandler.handle_error(e, "更新统计")

    def _update_stats_display(self, stats):
        """更新统计显示"""
        try:
            if hasattr(self, 'stats_label'):
                self.stats_label.config(
                    text=f"今日翻译: {stats['daily_translations']}次 ({stats['daily_characters']}字) | "
                        f"总计: {stats['total_translations']}次 ({stats['total_characters']}字)"
                )
        except Exception as e:
            logging.error(f"更新统计显示失败: {str(e)}")
    def export_translation(self):
        """导出翻译结果"""
        try:
            source_text = self.translate_tab_manager.source_text.get("1.0", "end").strip()
            target_text = self.translate_tab_manager.target_text.get("1.0", "end").strip()
            
            if not source_text or not target_text:
                Messagebox.show_warning("警告", "没有可导出的内容")
                return
                
            # 创建导出选项窗口
            export_window = tb.Toplevel(self.root)
            export_window.title("导出翻译")
            export_window.geometry("300x200")
            export_window.transient(self.root)
            export_window.grab_set()
            
            # 创建格式选择
            format_frame = tb.LabelFrame(export_window, text="选择导出格式", padding=10)
            format_frame.pack(fill=X, padx=20, pady=10)
            
            format_var = tb.StringVar(value="txt")
            formats = [
                ("纯文本 (.txt)", "txt"),
                ("Word文档 (.docx)", "docx"),
                ("PDF文档 (.pdf)", "pdf"),
                ("JSON (.json)", "json")
            ]
            
            for text, value in formats:
                tb.Radiobutton(format_frame, text=text, variable=format_var, 
                            value=value).pack(anchor="w", pady=2)
            
            # 添加确认按钮
            def do_export():
                format = format_var.get()
                if format == "txt":
                    self._export_txt(source_text, target_text)
                elif format == "docx":
                    self._export_docx(source_text, target_text)
                elif format == "pdf":
                    self._export_pdf(source_text, target_text)
                elif format == "json":
                    self._export_json(source_text, target_text)
                export_window.destroy()
                
            tb.Button(export_window, text="导出", bootstyle=SUCCESS,
                    command=do_export).pack(pady=10)
                    
        except Exception as e:
            logging.error(f"导出失败: {str(e)}")
            Messagebox.show_error("错误", f"导出失败: {str(e)}")

    def _export_txt(self, source_text, target_text):
        """导出为纯文本"""
        try:
            from tkinter import filedialog
            file_path = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[("Text files", "*.txt"), ("All files", "*.*")]
            )
            if file_path:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(f"原文：\n{source_text}\n\n译文：\n{target_text}")
                Messagebox.show_info("成功", "导出成功")
        except Exception as e:
            logging.error(f"导出文本失败: {str(e)}")
            Messagebox.show_error("错误", f"导出文本失败: {str(e)}")

    def _export_docx(self, source_text, target_text):
        """导出为Word文档"""
        try:
            from docx import Document  # 改为按需导入
            from tkinter import filedialog
            import sys
            
            # 处理打包后的路径问题
            if getattr(sys, 'frozen', False):
                base_path = sys._MEIPASS
            else:
                base_path = os.path.dirname(os.path.abspath(__file__))
                
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
            )
            if file_path:
                doc = Document()
                doc.add_heading('翻译结果', 0)
                doc.add_heading('原文', level=1)
                doc.add_paragraph(source_text)
                doc.add_heading('译文', level=1)
                doc.add_paragraph(target_text)
                doc.save(file_path)
                Messagebox.show_info("成功", "导出成功")
        except ImportError as e:
            Messagebox.show_error("错误", "导出Word文档需要安装python-docx库")
        except Exception as e:
            logging.error(f"导出Word文档失败: {str(e)}")
            Messagebox.show_error("错误", f"导出Word文档失败: {str(e)}")

    def _export_pdf(self, source_text, target_text):
        """导出为PDF文档"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from tkinter import filedialog
            
            file_path = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
            )
            if file_path:
                c = canvas.Canvas(file_path, pagesize=letter)
                # 注册中文字体
                try:
                    pdfmetrics.registerFont(TTFont('SimSun', 'SimSun.ttf'))
                    text_font = 'SimSun'
                except:
                    text_font = 'Helvetica'
                
                # 添加原文
                c.setFont(text_font, 12)
                c.drawString(100, 750, "原文：")
                y = 730
                for line in source_text.split('\n'):
                    c.drawString(100, y, line)
                    y -= 20
                    if y < 50:
                        c.showPage()
                        y = 750
                
                # 添加译文
                y -= 40
                c.drawString(100, y, "译文：")
                y -= 20
                for line in target_text.split('\n'):
                    c.drawString(100, y, line)
                    y -= 20
                    if y < 50:
                        c.showPage()
                        y = 750
                
                c.save()
                Messagebox.show_info("成功", "导出成功")
        except Exception as e:
            logging.error(f"导出PDF失败: {str(e)}")
            Messagebox.show_error("错误", f"导出PDF失败: {str(e)}")

    def _export_json(self, source_text, target_text):
        """导出为JSON"""
        try:
            import json
            from tkinter import filedialog
            
            file_path = filedialog.asksaveasfilename(
                defaultextension=".json",
                filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
            )
            if file_path:
                data = {
                    "source_text": source_text,
                    "target_text": target_text,
                    "timestamp": time.strftime('%Y-%m-%d %H:%M:%S')
                }
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                Messagebox.show_info("成功", "导出成功")
        except Exception as e:
            logging.error(f"导出JSON失败: {str(e)}")
            Messagebox.show_error("错误", f"导出JSON失败: {str(e)}")
class Tooltip:
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tooltip = None
        self.widget.bind("<Enter>", self.on_enter)
        self.widget.bind("<Leave>", self.on_leave)

    def on_enter(self, event=None):
        x, y, _, _ = self.widget.bbox("insert")
        x += self.widget.winfo_rootx() + 25
        y += self.widget.winfo_rooty() + 25
        self.tooltip = tb.Toplevel(self.widget)
        self.tooltip.wm_overrideredirect(True)
        self.tooltip.wm_geometry(f"+{x}+{y}")
        label = tb.Label(self.tooltip, text=self.text, 
                        bootstyle=INFO, padding=5)
        label.pack()

    def on_leave(self, event=None):
        if self.tooltip:
            self.tooltip.destroy()
            self.tooltip = None